from __future__ import annotations

import base64
import io
import inspect
import json
import logging
import re
import unicodedata
import uuid
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path
from time import perf_counter
from typing import Any, Awaitable, Callable, Dict, Iterable, List, Optional, Tuple
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from fastapi import UploadFile
import httpx
from pypdf import PdfReader
from pypdf.errors import PdfReadError
from sqlalchemy import desc
from sqlmodel import select
from sqlmodel.ext.asyncio.session import AsyncSession

from ...config import load_settings
from ...constants import MAX_LIMIT
from ...models import ExperienceCategory, ExperienceVersion, MasterExperience
from ..ai.ai_service import call_llm_json
from .prompts import RESUME_CHUNK_PARSING_PROMPT, RESUME_MERGE_PROMPT, RESUME_PARSING_PROMPT
from .schemas import DuplicateMatch, ParsedExperienceItem, ParsedExperienceVersion

logger = logging.getLogger(__name__)
settings = load_settings()

SUPPORTED_PDF_TYPES = {"application/pdf"}
SUPPORTED_DOCX_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
MIN_TEXT_LENGTH = 30
MIN_MEANINGFUL_TEXT_CHARS = 10
MAX_RESUME_TEXT_CHARS = 12_000
LONG_RESUME_TEXT_THRESHOLD = 6_000
CHUNK_MAX_CHARS = 3_200
CHUNK_MIN_CHARS = 800
MAX_MERGE_PAYLOAD_CHARS = 10_000
DUPLICATE_SIMILARITY_THRESHOLD = 0.86
DEFAULT_WORK_TITLE = "未命名经历"
DEFAULT_WORK_ORG = "未知机构"
DEFAULT_EDU_TITLE = "未命名专业"
DEFAULT_EDU_ORG = "未命名学校"
DEFAULT_SKILL_CATEGORY = "未分类"
PRESENT_MARKERS = {"present", "current", "now", "至今", "目前"}
COURSE_SPLIT_PATTERN = re.compile(r"[,，;；/\n]")
SKILL_TAG_SPLIT_PATTERN = re.compile(r"[,，;；/\n、|]+")
CJK_CHAR_PATTERN = r"\u4e00-\u9fff\u3400-\u4dbf"
CJK_PUNCT_PATTERN = r"\u3000-\u303f\uff00-\uffef·•"
CJK_INLINE_PATTERN = f"{CJK_CHAR_PATTERN}{CJK_PUNCT_PATTERN}"
CJK_PUNCT_ADJACENT_PATTERN = r"\(\)\[\]（）【】《》<>·•"
WHITESPACE_PATTERN = re.compile(r"\s+")
PARA_SPLIT_PATTERN = re.compile(r"\n\s*\n+")
SENTENCE_SPLIT_PATTERN = re.compile(r"(?<=[。！？!?;；.])\s+")
LINK_SPLIT_PATTERN = re.compile(r"[\s,;，；]+")
PERSONAL_INFO_FIELDS = ("full_name", "email", "phone", "location")
UNREADABLE_RESUME_TEXT_ERROR = (
    "无法读取附件中的文本内容，请检查上传内容；当前不支持无法读取文本的附件。"
)
PROJECT_KEYWORDS = {
    "project",
    "projects",
    "side project",
    "personal project",
    "open source",
    "opensource",
    "github",
    "开源",
    "项目",
    "课程设计",
    "课程项目",
    "毕业设计",
    "竞赛",
    "比赛",
    "作品",
}
WORK_KEYWORDS = {
    "intern",
    "internship",
    "full-time",
    "part-time",
    "employment",
    "company",
    "client",
    "客户",
    "公司",
    "集团",
    "部门",
    "岗位",
    "实习",
    "任职",
}
WORK_ORG_HINTS = {
    "有限公司",
    "股份有限公司",
    "有限责任公司",
    "公司",
    "集团",
    "inc",
    "ltd",
    "llc",
    "corp",
}
PROJECT_TITLE_HINTS = {"项目", "project"}
PROJECT_NAME_HINTS = {
    "system",
    "platform",
    "project",
    "app",
    "website",
    "web",
    "service",
    "tool",
    "dashboard",
    "系统",
    "平台",
    "项目",
    "应用",
    "网站",
    "小程序",
    "服务",
    "工具",
    "后台",
    "管理后台",
    "商城",
    "门户",
    "客户端",
    "gis",
    "webgis",
}
PROJECT_ROLE_HINTS = {
    "owner",
    "lead",
    "leader",
    "pm",
    "project manager",
    "tech lead",
    "engineer",
    "developer",
    "maintainer",
    "contributor",
    "负责人",
    "项目经理",
    "组长",
    "组员",
    "成员",
    "主导",
    "牵头",
    "独立开发",
    "核心开发",
    "开发",
}
PROJECT_MIN_SCORE = 1
LOG_WARN_THRESHOLDS_MS = {
    "read_file": 3_000,
    "parse_pdf": 8_000,
    "parse_docx": 5_000,
    "encode_file": 8_000,
    "ai_call": 20_000,
    "parse_resume_total": 25_000,
}
MAX_ATTACHMENT_BYTES = 5 * 1024 * 1024
ParseProgressCallback = Optional[Callable[[Dict[str, Any]], Awaitable[None] | None]]
ThoughtCallback = Optional[Callable[[Dict[str, Any]], Awaitable[None] | None]]
GEMINI_CONNECT_TIMEOUT_SECONDS = 10.0
GEMINI_POOL_TIMEOUT_SECONDS = 10.0


@dataclass(frozen=True)
class ExistingExperience:
    category: ExperienceCategory
    title: str
    org: str


def _normalize_text(value: Optional[str]) -> str:
    if not value:
        return ""
    compact = WHITESPACE_PATTERN.sub(" ", value).strip().lower()
    compact = re.sub(
        rf"(?<=[{CJK_INLINE_PATTERN}])\s+(?=[{CJK_INLINE_PATTERN}])",
        "",
        compact,
    )
    compact = re.sub(
        rf"(?<=[{CJK_CHAR_PATTERN}])\s+(?=[{CJK_PUNCT_ADJACENT_PATTERN}])",
        "",
        compact,
    )
    compact = re.sub(
        rf"(?<=[{CJK_PUNCT_ADJACENT_PATTERN}])\s+(?=[{CJK_CHAR_PATTERN}])",
        "",
        compact,
    )
    return compact


def _split_into_paragraphs(text: str) -> List[str]:
    stripped = text.strip()
    if not stripped:
        return []
    parts = [
        part.strip()
        for part in PARA_SPLIT_PATTERN.split(stripped)
        if part.strip()
    ]
    return parts if parts else [stripped]


def _hard_split_text(text: str, max_chars: int) -> List[str]:
    chunks: List[str] = []
    for index in range(0, len(text), max_chars):
        piece = text[index : index + max_chars].strip()
        if piece:
            chunks.append(piece)
    return chunks


def _chunk_units(units: Iterable[str], joiner: str) -> List[str]:
    chunks: List[str] = []
    current: List[str] = []
    current_len = 0
    joiner_len = len(joiner)
    for unit in units:
        cleaned = unit.strip()
        if not cleaned:
            continue
        if len(cleaned) > CHUNK_MAX_CHARS:
            if current:
                chunks.append(joiner.join(current).strip())
                current = []
                current_len = 0
            chunks.extend(_hard_split_text(cleaned, CHUNK_MAX_CHARS))
            continue
        projected = current_len + len(cleaned) + (joiner_len if current else 0)
        if projected <= CHUNK_MAX_CHARS:
            current.append(cleaned)
            current_len = projected
            continue
        if current:
            chunks.append(joiner.join(current).strip())
        current = [cleaned]
        current_len = len(cleaned)
    if current:
        chunks.append(joiner.join(current).strip())
    return chunks


def _split_long_paragraph(paragraph: str) -> List[str]:
    if len(paragraph) <= CHUNK_MAX_CHARS:
        return [paragraph]
    lines = [line.strip() for line in paragraph.splitlines() if line.strip()]
    if len(lines) > 1:
        line_chunks = _chunk_units(lines, "\n")
        if line_chunks:
            return line_chunks
    sentences = [
        item.strip()
        for item in SENTENCE_SPLIT_PATTERN.split(paragraph)
        if item.strip()
    ]
    if len(sentences) > 1:
        return _chunk_units(sentences, " ")
    return _hard_split_text(paragraph, CHUNK_MAX_CHARS)


def _chunk_paragraphs(paragraphs: Iterable[str]) -> List[str]:
    chunks: List[str] = []
    current: List[str] = []
    current_len = 0
    for paragraph in paragraphs:
        parts = _split_long_paragraph(paragraph)
        for part in parts:
            projected = current_len + len(part) + (1 if current else 0)
            if projected <= CHUNK_MAX_CHARS:
                current.append(part)
                current_len = projected
                continue
            if current:
                chunks.append("\n".join(current).strip())
            current = [part]
            current_len = len(part)
    if current:
        chunks.append("\n".join(current).strip())
    return chunks


def _merge_small_chunks(chunks: List[str]) -> List[str]:
    if not chunks:
        return []
    merged: List[str] = []
    buffer = ""
    for chunk in chunks:
        if not buffer:
            buffer = chunk
            continue
        if (
            len(buffer) < CHUNK_MIN_CHARS
            and len(buffer) + len(chunk) + 1 <= CHUNK_MAX_CHARS
        ):
            buffer = f"{buffer}\n{chunk}".strip()
            continue
        merged.append(buffer)
        buffer = chunk
    if buffer:
        merged.append(buffer)
    return merged


def _split_resume_text(text: str) -> List[str]:
    paragraphs = _split_into_paragraphs(text)
    if not paragraphs:
        return [text] if text.strip() else []
    chunks = _chunk_paragraphs(paragraphs)
    chunks = [chunk for chunk in _merge_small_chunks(chunks) if chunk]
    return chunks or [text]


def _should_use_chunking(text: str) -> bool:
    return len(text) > LONG_RESUME_TEXT_THRESHOLD


def _build_signature(title: Optional[str], org: Optional[str]) -> str:
    title_part = _normalize_text(title)
    org_part = _normalize_text(org)
    if not title_part and not org_part:
        return ""
    return f"{org_part}::{title_part}"


def _similarity(a: str, b: str) -> float:
    if not a or not b:
        return 0.0
    return SequenceMatcher(None, a, b).ratio()


def _is_present_marker(value: Any) -> bool:
    if not value:
        return False
    marker = str(value).strip().lower()
    return marker in PRESENT_MARKERS


def _normalize_date(value: Any) -> Optional[str]:
    if not value:
        return None
    text = str(value).strip()
    if not text or _is_present_marker(text):
        return None
    if re.match(r"^\d{4}-\d{2}-\d{2}$", text):
        return text
    if re.match(r"^\d{4}-\d{2}$", text):
        return f"{text}-01"
    if re.match(r"^\d{4}$", text):
        return f"{text}-01-01"
    return None


def _ensure_list(value: Any) -> List[Any]:
    if isinstance(value, list):
        return value
    return []


def _ensure_str(value: Any) -> str:
    if value is None:
        return ""
    return _clean_inline_text(str(value))


def _clean_inline_text(value: str) -> str:
    compact = WHITESPACE_PATTERN.sub(" ", value).strip()
    compact = re.sub(
        rf"(?<=[{CJK_INLINE_PATTERN}])\s+(?=[{CJK_INLINE_PATTERN}])",
        "",
        compact,
    )
    compact = re.sub(
        rf"(?<=[{CJK_CHAR_PATTERN}])\s+(?=[{CJK_PUNCT_ADJACENT_PATTERN}])",
        "",
        compact,
    )
    compact = re.sub(
        rf"(?<=[{CJK_PUNCT_ADJACENT_PATTERN}])\s+(?=[{CJK_CHAR_PATTERN}])",
        "",
        compact,
    )
    return compact


def _clean_multiline_text(value: str) -> str:
    lines = [line for line in value.splitlines()]
    cleaned = [_clean_inline_text(line) for line in lines]
    return "\n".join([line for line in cleaned if line])


def _clean_resume_text(text: str) -> str:
    if not text:
        return text
    lines = text.splitlines()
    cleaned_lines = [
        _clean_inline_text(line) if line.strip() else ""
        for line in lines
    ]
    return "\n".join(cleaned_lines)


def _ensure_optional_str(value: Any) -> Optional[str]:
    text = _ensure_str(value)
    return text or None


def _normalize_text_list(value: Any) -> List[str]:
    return [
        _clean_inline_text(item)
        for item in _ensure_list(value)
        if isinstance(item, str) and item.strip()
    ]


def _normalize_skill_tags(value: Any) -> List[str]:
    if isinstance(value, list):
        return _normalize_text_list(value)
    if isinstance(value, str):
        items = [
            item.strip()
            for item in SKILL_TAG_SPLIT_PATTERN.split(value)
            if item.strip()
        ]
        return items
    return []


def _join_text_parts(parts: Iterable[str]) -> str:
    return " ".join([part for part in parts if part]).strip()


def _contains_any(text: str, keywords: Iterable[str]) -> bool:
    if not text:
        return False
    return any(keyword in text for keyword in keywords)


def _count_keyword_hits(text: str, keywords: Iterable[str]) -> int:
    if not text:
        return 0
    return sum(1 for keyword in keywords if keyword in text)


def _collect_entry_text(entry: Dict[str, Any]) -> str:
    star_source = entry.get("star") if isinstance(entry.get("star"), dict) else {}
    parts = [
        _ensure_str(entry.get("title")),
        _ensure_str(entry.get("org")),
        _ensure_str(entry.get("summary")),
        " ".join(_normalize_text_list(entry.get("highlights"))),
        _ensure_str(star_source.get("s")),
        _ensure_str(star_source.get("t")),
        _ensure_str(star_source.get("a")),
        _ensure_str(star_source.get("r")),
    ]
    return _join_text_parts(parts)


def _normalize_courses(value: Any) -> Any:
    if isinstance(value, list):
        return _normalize_text_list(value)
    if isinstance(value, str):
        items = [
            _clean_inline_text(item)
            for item in COURSE_SPLIT_PATTERN.split(value)
            if item.strip()
        ]
        return items if items else _clean_inline_text(value)
    return ""


def _normalize_star_field(value: Any) -> str:
    if isinstance(value, list):
        items = [
            _clean_inline_text(item)
            for item in value
            if isinstance(item, str) and item.strip()
        ]
        return "\n".join(items)
    if value is None:
        return ""
    return _clean_multiline_text(str(value))


def _resolve_action_text(current: str, highlights: Any) -> str:
    highlight_items = _normalize_text_list(highlights)
    if not highlight_items:
        return current
    highlight_text = "\n".join(highlight_items)
    if not current or len(highlight_text) > len(current):
        return highlight_text
    return current


def _build_star_payload(entry: Dict[str, Any]) -> Dict[str, str]:
    star_source = entry.get("star") if isinstance(entry.get("star"), dict) else {}
    s_value = _normalize_star_field(star_source.get("s") or entry.get("s"))
    t_value = _normalize_star_field(star_source.get("t") or entry.get("t"))
    a_value = _normalize_star_field(star_source.get("a") or entry.get("a"))
    r_value = _normalize_star_field(star_source.get("r") or entry.get("r"))
    a_value = _resolve_action_text(a_value, entry.get("highlights"))
    return {"s": s_value, "t": t_value, "a": a_value, "r": r_value}


def _build_work_version(entry: Dict[str, Any]) -> ParsedExperienceVersion:
    title = _ensure_str(entry.get("title")) or DEFAULT_WORK_TITLE
    org = _ensure_str(entry.get("org")) or DEFAULT_WORK_ORG
    end_raw = entry.get("end_date")
    is_current = bool(entry.get("is_current")) or _is_present_marker(end_raw)
    return ParsedExperienceVersion(
        title=title,
        org=org,
        location=_ensure_str(entry.get("location")) or None,
        start_date=_normalize_date(entry.get("start_date")),
        end_date=_normalize_date(end_raw),
        is_current=is_current,
        summary=_ensure_str(entry.get("summary")) or None,
        highlights=_normalize_text_list(entry.get("highlights")),
        tags=_normalize_text_list(entry.get("tags")),
        star=_build_star_payload(entry),
    )


def _should_swap_project_fields(title: str, org: str) -> bool:
    title_text = _normalize_text(title)
    org_text = _normalize_text(org)
    if not title_text or not org_text:
        return False
    title_project = _count_keyword_hits(title_text, PROJECT_NAME_HINTS)
    org_project = _count_keyword_hits(org_text, PROJECT_NAME_HINTS)
    title_role = _count_keyword_hits(title_text, PROJECT_ROLE_HINTS)
    org_role = _count_keyword_hits(org_text, PROJECT_ROLE_HINTS)
    swap_score = title_project + org_role
    keep_score = org_project + title_role
    return swap_score > keep_score and swap_score >= PROJECT_MIN_SCORE


def _normalize_project_entry(entry: Dict[str, Any]) -> Dict[str, Any]:
    title = _ensure_str(entry.get("title"))
    org = _ensure_str(entry.get("org"))
    if _should_swap_project_fields(title, org):
        return {**entry, "title": org, "org": title}
    return entry


def _build_project_version(entry: Dict[str, Any]) -> ParsedExperienceVersion:
    normalized = _normalize_project_entry(entry)
    return _build_work_version(normalized)


def _build_education_version(entry: Dict[str, Any]) -> ParsedExperienceVersion:
    school = _ensure_str(entry.get("school"))
    major = _ensure_str(entry.get("major"))
    degree = _ensure_str(entry.get("degree"))
    title = major or degree or DEFAULT_EDU_TITLE
    org = school or DEFAULT_EDU_ORG
    star: Dict[str, Any] = {}
    if degree:
        star["degree"] = degree
    gpa = _ensure_str(entry.get("gpa"))
    if gpa:
        star["gpa"] = gpa
    courses = _normalize_courses(entry.get("courses"))
    if courses:
        star["courses"] = courses
    is_current = bool(entry.get("is_current")) or _is_present_marker(entry.get("end_date"))
    return ParsedExperienceVersion(
        title=title,
        org=org,
        start_date=_normalize_date(entry.get("start_date")),
        end_date=_normalize_date(entry.get("end_date")),
        is_current=is_current,
        star=star,
    )


def _build_item(category: ExperienceCategory, version: ParsedExperienceVersion) -> ParsedExperienceItem:
    return ParsedExperienceItem(
        id=str(uuid.uuid4()),
        category=category,
        version=version,
    )


def _infer_experience_category(entry: Dict[str, Any]) -> ExperienceCategory:
    text = _normalize_text(_collect_entry_text(entry))
    title = _normalize_text(_ensure_str(entry.get("title")))
    org = _normalize_text(_ensure_str(entry.get("org")))
    project_score = _count_keyword_hits(text, PROJECT_KEYWORDS)
    work_score = _count_keyword_hits(text, WORK_KEYWORDS)

    if _contains_any(org, WORK_ORG_HINTS):
        work_score += 1
    if _contains_any(title, PROJECT_TITLE_HINTS):
        project_score += 1

    if project_score >= PROJECT_MIN_SCORE and project_score > work_score:
        return ExperienceCategory.PROJECT
    return ExperienceCategory.WORK


def _split_work_and_project_entries(
    entries: Iterable[Any],
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    work_entries: List[Dict[str, Any]] = []
    project_entries: List[Dict[str, Any]] = []
    for entry in entries:
        if not isinstance(entry, dict):
            continue
        category = _infer_experience_category(entry)
        if category == ExperienceCategory.PROJECT:
            project_entries.append(entry)
        else:
            work_entries.append(entry)
    return work_entries, project_entries


def build_resume_items(payload: Dict[str, Any]) -> List[ParsedExperienceItem]:
    items: List[ParsedExperienceItem] = []
    raw_work_entries = _ensure_list(payload.get("work_experiences"))
    project_payload = (
        payload.get("project_experiences")
        if "project_experiences" in payload
        else None
    )
    raw_project_entries = _ensure_list(project_payload)
    if "project_experiences" in payload:
        work_entries = [entry for entry in raw_work_entries if isinstance(entry, dict)]
        project_entries = [entry for entry in raw_project_entries if isinstance(entry, dict)]
    else:
        work_entries, project_entries = _split_work_and_project_entries(raw_work_entries)

    for entry in work_entries:
        items.append(_build_item(ExperienceCategory.WORK, _build_work_version(entry)))
    for entry in project_entries:
        items.append(_build_item(ExperienceCategory.PROJECT, _build_project_version(entry)))
    for entry in _ensure_list(payload.get("education")):
        if isinstance(entry, dict):
            items.append(
                _build_item(ExperienceCategory.EDUCATION, _build_education_version(entry))
            )
    return items


def _resolve_file_kind(file: UploadFile) -> str:
    filename = file.filename or ""
    extension = Path(filename).suffix.lower()
    content_type = (file.content_type or "").lower()
    if content_type in SUPPORTED_PDF_TYPES or extension == ".pdf":
        return "pdf"
    if content_type in SUPPORTED_DOCX_TYPES or extension == ".docx":
        return "docx"
    raise ValueError("不支持的文件类型，请上传 PDF 或 DOCX 文件。")


def _ensure_attachment_size_limit(data: bytes) -> None:
    if len(data) <= MAX_ATTACHMENT_BYTES:
        return
    max_mb = MAX_ATTACHMENT_BYTES / (1024 * 1024)
    raise ValueError(
        f"文件过大，无法直接解析。请上传不超过 {max_mb:.0f}MB 的 PDF 或 DOCX 文件。"
    )


def _resolve_file_mime(file: UploadFile, kind: str) -> str:
    content_type = (file.content_type or "").lower().strip()
    if content_type:
        return content_type
    if kind == "pdf":
        return "application/pdf"
    return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _resolve_file_kind_from_metadata(filename: str, content_type: str) -> str:
    extension = Path(filename or "").suffix.lower()
    normalized_type = (content_type or "").lower().strip()
    if normalized_type in SUPPORTED_PDF_TYPES or extension == ".pdf":
        return "pdf"
    if normalized_type in SUPPORTED_DOCX_TYPES or extension == ".docx":
        return "docx"
    raise ValueError("不支持的文件类型，请上传 PDF 或 DOCX 文件。")


def _log_timing(
    step: str,
    duration_ms: float,
    request_id: Optional[str],
    extra: Optional[Dict[str, Any]] = None,
) -> None:
    payload = {"step": step, "duration_ms": round(duration_ms, 2)}
    if request_id:
        payload["request_id"] = request_id
    if extra:
        payload.update(extra)
    threshold = LOG_WARN_THRESHOLDS_MS.get(step)
    if threshold is not None and duration_ms >= threshold:
        logger.warning("[ResumeParse] %s", payload)
        return
    logger.info("[ResumeParse] %s", payload)


def _build_resume_messages(prompt: str, content: str) -> List[Dict[str, Any]]:
    return [
        {"role": "system", "content": prompt},
        {"role": "user", "content": content},
    ]


def _encode_upload_data(data: bytes) -> str:
    return base64.b64encode(data).decode("ascii")


def _build_resume_attachment_messages(
    prompt: str,
    filename: str,
    mime_type: str,
    encoded_file: str,
) -> List[Dict[str, Any]]:
    data_url = f"data:{mime_type};base64,{encoded_file}"
    return [
        {"role": "system", "content": prompt},
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": (
                        "请直接解析附件简历，并严格按 JSON schema 返回。"
                        f"文件名：{filename or 'resume'}"
                    ),
                },
                {
                    "type": "file_url",
                    "file_url": {"url": data_url},
                },
            ],
        },
    ]


async def _emit_progress(
    progress_callback: ParseProgressCallback,
    payload: Dict[str, Any],
) -> None:
    if not progress_callback:
        return
    result = progress_callback(payload)
    if inspect.isawaitable(result):
        await result


async def _emit_thought(
    thought_callback: ThoughtCallback,
    payload: Dict[str, Any],
) -> None:
    if not thought_callback:
        return
    result = thought_callback(payload)
    if inspect.isawaitable(result):
        await result


def _build_gemini_headers() -> Dict[str, str]:
    api_key = settings.gemini_api_key
    if not api_key:
        raise ValueError("GEMINI_API_KEY 未配置，无法返回 Gemini 实时思考节点。")
    return {
        "x-goog-api-key": api_key,
        "Content-Type": "application/json",
    }


def _build_gemini_stream_url(model: str) -> str:
    base_url = (settings.gemini_base_url or "").rstrip("/")
    if not base_url:
        raise ValueError("GEMINI_BASE_URL 未配置，无法调用 Gemini Thinking。")
    normalized = base_url.lower()
    if not normalized.endswith("/v1beta") and not normalized.endswith("/v1"):
        base_url = f"{base_url}/v1beta"
    return f"{base_url}/models/{model}:streamGenerateContent?alt=sse"


def _build_gemini_timeout() -> httpx.Timeout:
    return httpx.Timeout(
        connect=GEMINI_CONNECT_TIMEOUT_SECONDS,
        write=float(settings.ai_timeout_seconds),
        read=float(settings.ai_timeout_seconds),
        pool=GEMINI_POOL_TIMEOUT_SECONDS,
    )


def _build_resume_thinking_request(cleaned_text: str) -> Dict[str, Any]:
    return {
        "systemInstruction": {
            "parts": [{"text": RESUME_PARSING_PROMPT}],
        },
        "contents": [
            {
                "role": "user",
                "parts": [
                    {
                        "text": (
                            "请解析以下简历正文，并严格输出 JSON。"
                            "不要补充正文中不存在的信息。\n\n"
                            f"{cleaned_text}"
                        )
                    }
                ],
            }
        ],
        "generationConfig": {
            "temperature": 0.2,
            "responseMimeType": "application/json",
            "thinkingConfig": {
                "includeThoughts": True,
            },
        },
    }


async def _iter_sse_json_payloads(response: httpx.Response):
    def build_payload(lines: List[str]) -> str:
        data_lines: List[str] = []
        for item in lines:
            if not item.startswith("data:"):
                continue
            value = item[5:]
            if value.startswith(" "):
                value = value[1:]
            data_lines.append(value)
        return "\n".join(data_lines)

    event_lines: List[str] = []
    async for raw_line in response.aiter_lines():
        line = raw_line.rstrip("\r")
        if not line.strip():
            if not event_lines:
                continue
            payload = build_payload(event_lines)
            event_lines = []
            if not payload:
                continue
            if payload == "[DONE]":
                break
            try:
                yield json.loads(payload)
            except json.JSONDecodeError:
                logger.warning("[ResumeParse] invalid Gemini SSE payload: %s", payload[:500])
            continue
        event_lines.append(line)

    if event_lines:
        payload = build_payload(event_lines)
        if payload and payload != "[DONE]":
            try:
                yield json.loads(payload)
            except json.JSONDecodeError:
                logger.warning("[ResumeParse] invalid Gemini SSE trailing payload: %s", payload[:500])


async def _stream_resume_thinking_parse(
    cleaned_text: str,
    request_id: Optional[str],
    thought_callback: ThoughtCallback = None,
) -> Dict[str, Any]:
    request_body = _build_resume_thinking_request(cleaned_text)
    model = settings.gemini_model
    url = _build_gemini_stream_url(model)
    answer_parts: List[str] = []
    call_start = perf_counter()

    try:
        async with httpx.AsyncClient(timeout=_build_gemini_timeout()) as client:
            async with client.stream(
                "POST",
                url,
                headers=_build_gemini_headers(),
                json=request_body,
            ) as response:
                response.raise_for_status()
                content_type = (response.headers.get("content-type") or "").lower()
                if "text/event-stream" not in content_type:
                    body_preview = (await response.aread()).decode("utf-8", errors="ignore")[:800]
                    logger.error(
                        "[ResumeParse] Gemini proxy returned unexpected content-type request_id=%s content_type=%s body=%s",
                        request_id,
                        content_type,
                        body_preview,
                    )
                    raise ValueError(
                        "Gemini 中转站返回了非流式响应，请检查 GEMINI_BASE_URL 是否需要包含 /v1beta。"
                    )
                async for payload in _iter_sse_json_payloads(response):
                    candidates = payload.get("candidates") or []
                    if not candidates:
                        continue
                    parts = ((candidates[0] or {}).get("content") or {}).get("parts") or []
                    for part in parts:
                        text = part.get("text")
                        if not isinstance(text, str) or not text:
                            continue
                        if part.get("thought") is True:
                            await _emit_thought(
                                thought_callback,
                                {"type": "thought", "summary": text},
                            )
                            continue
                        answer_parts.append(text)
    except httpx.HTTPStatusError as exc:
        try:
            await exc.response.aread()
            error_text = exc.response.text[:1000]
        except Exception:
            error_text = "Failed to read response body."
        logger.error(
            "[ResumeParse] Gemini thinking request failed request_id=%s status=%s body=%s",
            request_id,
            exc.response.status_code,
            error_text,
        )
        raise ValueError("Gemini Thinking 解析失败，请稍后重试。") from exc
    except httpx.TimeoutException as exc:
        raise ValueError("Gemini Thinking 解析超时，请稍后重试。") from exc

    call_ms = (perf_counter() - call_start) * 1000
    _log_timing(
        "ai_call",
        call_ms,
        request_id,
        {
            "mode": "gemini_thinking",
            "input_length": len(cleaned_text),
        },
    )

    answer_text = "".join(answer_parts).strip()
    if not answer_text:
        raise ValueError("Gemini 未返回可解析的结构化结果。")
    return _normalize_parse_result(_parse_structured_response_text(answer_text))


async def _call_resume_llm(
    messages: List[Dict[str, Any]],
    request_id: Optional[str],
    step: str,
    extra: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    call_start = perf_counter()
    try:
        result = await call_llm_json(messages)
    except Exception as exc:
        call_ms = (perf_counter() - call_start) * 1000
        payload = {"status": "error", "error": type(exc).__name__}
        if extra:
            payload.update(extra)
        _log_timing(step, call_ms, request_id, payload)
        raise
    call_ms = (perf_counter() - call_start) * 1000
    payload = {"status": "ok"}
    if extra:
        payload.update(extra)
    _log_timing(step, call_ms, request_id, payload)
    return result


def _parse_structured_response_text(text: str) -> Dict[str, Any]:
    cleaned = text.strip()
    if not cleaned:
        raise ValueError("模型返回内容为空。")
    try:
        parsed = json.loads(cleaned)
    except json.JSONDecodeError:
        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start < 0 or end <= start:
            raise ValueError("模型返回的结构化内容不是合法 JSON。")
        try:
            parsed = json.loads(cleaned[start : end + 1])
        except json.JSONDecodeError as exc:
            raise ValueError("模型返回的结构化内容不是合法 JSON。") from exc
    if not isinstance(parsed, dict):
        raise ValueError("模型返回的结构化内容格式不正确。")
    return parsed


def _normalize_parse_result(result: Any) -> Dict[str, Any]:
    if not isinstance(result, dict):
        raise ValueError("模型返回的结果格式不正确。")
    return result


def _extract_dict_entries(value: Any) -> List[Dict[str, Any]]:
    return [item for item in _ensure_list(value) if isinstance(item, dict)]


def _normalize_personal_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return _clean_inline_text(value)
    return _clean_inline_text(str(value))


def _normalize_personal_links(value: Any) -> List[str]:
    if not value:
        return []
    items: List[str] = []
    if isinstance(value, list):
        for item in value:
            if isinstance(item, str):
                items.extend(LINK_SPLIT_PATTERN.split(item))
    elif isinstance(value, str):
        items.extend(LINK_SPLIT_PATTERN.split(value))
    else:
        return []
    return [item.strip() for item in items if isinstance(item, str) and item.strip()]


def normalize_personal_info(payload: Any) -> Optional[Dict[str, Any]]:
    if not isinstance(payload, dict):
        return None
    personal_info = payload.get("personal_info")
    if not isinstance(personal_info, dict):
        return None
    normalized: Dict[str, Any] = {}
    for field in PERSONAL_INFO_FIELDS:
        value = _normalize_personal_value(personal_info.get(field))
        if value:
            normalized[field] = value
    links = _normalize_personal_links(personal_info.get("links"))
    if links:
        normalized["links"] = links
    return normalized or None


def _normalize_certification_entry(entry: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    name = _ensure_str(entry.get("name"))
    if not name:
        return None
    return {
        "name": name,
        "issuer": _ensure_optional_str(entry.get("issuer")),
        "issue_date": _ensure_optional_str(entry.get("issue_date")),
        "expiry_date": _ensure_optional_str(entry.get("expiry_date")),
        "credential_id": _ensure_optional_str(entry.get("credential_id")),
        "credential_url": _ensure_optional_str(entry.get("credential_url")),
        "description": _ensure_optional_str(entry.get("description")),
    }


def normalize_certifications(payload: Any) -> List[Dict[str, Any]]:
    if not isinstance(payload, dict):
        return []
    raw_items = _extract_dict_entries(payload.get("certifications"))
    normalized: List[Dict[str, Any]] = []
    for entry in raw_items:
        item = _normalize_certification_entry(entry)
        if item:
            normalized.append(item)
    if not normalized:
        return []
    return _dedupe_entries(normalized, ("name", "issuer", "issue_date"))


def _normalize_skill_category(value: Any) -> str:
    return _ensure_str(value) or DEFAULT_SKILL_CATEGORY


def normalize_skill_groups(payload: Any) -> List[Dict[str, Any]]:
    if not isinstance(payload, dict):
        return []
    raw_groups = _extract_dict_entries(payload.get("skills"))
    if not raw_groups:
        return []
    grouped: Dict[str, Dict[str, Any]] = {}
    order: List[str] = []
    for entry in raw_groups:
        category = _normalize_skill_category(entry.get("category"))
        category_key = _normalize_text(category)
        tags = _normalize_skill_tags(entry.get("tags"))
        if not tags:
            continue
        if category_key not in grouped:
            grouped[category_key] = {"category": category, "tags": []}
            order.append(category_key)
        existing_tags = grouped[category_key]["tags"]
        seen = { _normalize_text(tag) for tag in existing_tags }
        for tag in tags:
            tag_key = _normalize_text(tag)
            if not tag_key or tag_key in seen:
                continue
            existing_tags.append(tag)
            seen.add(tag_key)
    return [grouped[key] for key in order if grouped.get(key)]


def _merge_personal_info(results: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    merged: Dict[str, Any] = {}
    for field in PERSONAL_INFO_FIELDS:
        for result in results:
            personal_info = result.get("personal_info")
            if not isinstance(personal_info, dict):
                continue
            value = _normalize_personal_value(personal_info.get(field))
            if value:
                merged[field] = value
                break
    links: List[str] = []
    seen = set()
    for result in results:
        personal_info = result.get("personal_info")
        if not isinstance(personal_info, dict):
            continue
        for link in _normalize_personal_links(personal_info.get("links")):
            if link not in seen:
                seen.add(link)
                links.append(link)
    if links:
        merged["links"] = links
    return merged or None


def _entry_signature(entry: Dict[str, Any], fields: Tuple[str, ...]) -> str:
    parts = [
        _normalize_text(_ensure_str(entry.get(field)))
        for field in fields
    ]
    if not any(parts):
        return ""
    return "::".join(parts)


def _entry_score(entry: Dict[str, Any]) -> int:
    score = 0
    for value in entry.values():
        if isinstance(value, str):
            score += len(value.strip())
        elif isinstance(value, list):
            score += sum(len(item.strip()) for item in value if isinstance(item, str))
        elif isinstance(value, dict):
            score += sum(
                len(str(item).strip())
                for item in value.values()
                if isinstance(item, str)
            )
    return score


def _dedupe_entries(
    entries: List[Dict[str, Any]], fields: Tuple[str, ...]
) -> List[Dict[str, Any]]:
    output: List[Dict[str, Any]] = []
    index: Dict[str, int] = {}
    for entry in entries:
        signature = _entry_signature(entry, fields)
        if not signature:
            output.append(entry)
            continue
        existing_index = index.get(signature)
        if existing_index is None:
            index[signature] = len(output)
            output.append(entry)
            continue
        if _entry_score(entry) > _entry_score(output[existing_index]):
            output[existing_index] = entry
    return output


def _merge_chunk_results(results: List[Dict[str, Any]]) -> Dict[str, Any]:
    merged: Dict[str, Any] = {
        "work_experiences": [],
        "project_experiences": [],
        "education": [],
        "certifications": [],
        "skills": [],
    }
    personal_info = _merge_personal_info(results)
    if personal_info:
        merged["personal_info"] = personal_info
    for result in results:
        merged["work_experiences"].extend(
            _extract_dict_entries(result.get("work_experiences"))
        )
        merged["project_experiences"].extend(
            _extract_dict_entries(result.get("project_experiences"))
        )
        merged["education"].extend(
            _extract_dict_entries(result.get("education"))
        )
        merged["certifications"].extend(
            _extract_dict_entries(result.get("certifications"))
        )
        merged["skills"].extend(
            _extract_dict_entries(result.get("skills"))
        )
    merged["work_experiences"] = _dedupe_entries(
        merged["work_experiences"],
        ("title", "org", "start_date", "end_date"),
    )
    merged["project_experiences"] = _dedupe_entries(
        merged["project_experiences"],
        ("title", "org", "start_date", "end_date"),
    )
    merged["education"] = _dedupe_entries(
        merged["education"],
        ("school", "major", "degree", "start_date", "end_date"),
    )
    merged["certifications"] = normalize_certifications(merged)
    merged["skills"] = normalize_skill_groups(merged)
    return merged


async def _merge_with_llm(
    draft: Dict[str, Any], request_id: Optional[str]
) -> Dict[str, Any]:
    payload = json.dumps(draft, ensure_ascii=False)
    if len(payload) > MAX_MERGE_PAYLOAD_CHARS:
        return draft
    messages = _build_resume_messages(RESUME_MERGE_PROMPT, payload)
    try:
        result = await _call_resume_llm(
            messages,
            request_id,
            "ai_merge_call",
            {"input_length": len(payload)},
        )
    except Exception:
        return draft
    if not isinstance(result, dict):
        return draft
    return result


async def _parse_resume_single(
    text: str, request_id: Optional[str]
) -> Dict[str, Any]:
    trimmed = text
    if len(trimmed) > MAX_RESUME_TEXT_CHARS:
        trimmed = trimmed[:MAX_RESUME_TEXT_CHARS]
    messages = _build_resume_messages(RESUME_PARSING_PROMPT, trimmed)
    result = await _call_resume_llm(
        messages,
        request_id,
        "ai_call",
        {"input_length": len(trimmed)},
    )
    return _normalize_parse_result(result)


async def _parse_resume_chunked(
    text: str, request_id: Optional[str]
) -> Dict[str, Any]:
    chunks = _split_resume_text(text)
    chunk_results: List[Dict[str, Any]] = []
    for index, chunk in enumerate(chunks):
        messages = _build_resume_messages(RESUME_CHUNK_PARSING_PROMPT, chunk)
        try:
            result = await _call_resume_llm(
                messages,
                request_id,
                "ai_chunk_call",
                {
                    "chunk_index": index + 1,
                    "chunk_total": len(chunks),
                    "input_length": len(chunk),
                },
            )
        except Exception:
            continue
        if isinstance(result, dict):
            chunk_results.append(result)
    if not chunk_results:
        return await _parse_resume_single(text, request_id)
    draft = _merge_chunk_results(chunk_results)
    merged = await _merge_with_llm(draft, request_id)
    return _normalize_parse_result(merged)


async def extract_text(file: UploadFile, request_id: Optional[str] = None) -> bytes:
    total_start = perf_counter()
    read_start = perf_counter()
    data = await file.read()
    read_ms = (perf_counter() - read_start) * 1000
    _log_timing(
        "read_file",
        read_ms,
        request_id,
        {
            "size": len(data),
            "filename": file.filename or "",
            "content_type": file.content_type or "",
        },
    )
    if not data:
        raise ValueError("文件为空，无法解析。")
    _resolve_file_kind(file)
    total_ms = (perf_counter() - total_start) * 1000
    _log_timing("extract_text_total", total_ms, request_id)
    return data


def _extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def _extract_docx_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

    def append_table_text(table) -> None:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_parts = [para.text.strip() for para in cell.paragraphs if para.text.strip()]
                for nested_table in cell.tables:
                    append_table_text(nested_table)
                if cell_parts:
                    cells.append("\n".join(cell_parts))
            if cells:
                parts.append(" | ".join(cells))

    for table in doc.tables:
        append_table_text(table)

    return "\n".join(parts)


def extract_resume_text(
    file_data: bytes,
    filename: str,
    file_mime_type: str,
    request_id: Optional[str] = None,
) -> str:
    if not file_data:
        raise ValueError("文件为空，无法解析。")
    _ensure_attachment_size_limit(file_data)
    kind = _resolve_file_kind_from_metadata(filename, file_mime_type)
    parse_start = perf_counter()
    try:
        if kind == "pdf":
            text = _extract_pdf_text(file_data)
            parse_step = "parse_pdf"
        else:
            text = _extract_docx_text(file_data)
            parse_step = "parse_docx"
    except (PdfReadError, BadZipFile, PackageNotFoundError, ValueError) as exc:
        logger.warning(
            "[ResumeParse] failed to extract resume text request_id=%s filename=%s kind=%s error=%s",
            request_id,
            filename,
            kind,
            str(exc),
        )
        raise ValueError("文件无法读取，请确认文件未损坏、未加密且内容可解析。") from exc
    parse_ms = (perf_counter() - parse_start) * 1000
    _log_timing(parse_step, parse_ms, request_id, {"text_length": len(text)})
    return text


def _count_meaningful_text_chars(text: str) -> int:
    return sum(
        1
        for char in text
        if unicodedata.category(char).startswith(("L", "N"))
    )


def _prepare_resume_text(text: str, request_id: Optional[str] = None) -> Optional[str]:
    cleaned = _clean_resume_text(text)
    stripped = cleaned.strip()
    meaningful_char_count = _count_meaningful_text_chars(stripped)
    if meaningful_char_count < MIN_MEANINGFUL_TEXT_CHARS:
        logger.info(
            "[ResumeParse] extracted text unreadable, fallback to attachment request_id=%s text_length=%s meaningful_char_count=%s",
            request_id,
            len(stripped),
            meaningful_char_count,
        )
        return None
    if len(stripped) < MIN_TEXT_LENGTH:
        logger.info(
            "[ResumeParse] extracted text too short, fallback to attachment request_id=%s text_length=%s meaningful_char_count=%s",
            request_id,
            len(stripped),
            meaningful_char_count,
        )
        return None
    return cleaned


async def _parse_resume_from_text(
    *,
    cleaned_text: str,
    request_id: Optional[str],
    progress_callback: ParseProgressCallback = None,
) -> Dict[str, Any]:
    use_chunking = _should_use_chunking(cleaned_text)
    await _emit_progress(
        progress_callback,
        {
            "type": "progress",
            "node": "segment_resume",
            "title": "切分简历结构" if use_chunking else "构建解析上下文",
        },
    )
    await _emit_progress(
        progress_callback,
        {"type": "progress", "node": "request_ai", "title": "调用 AI 结构化解析"},
    )

    total_start = perf_counter()
    if use_chunking:
        result = await _parse_resume_chunked(cleaned_text, request_id)
        mode = "chunked"
    else:
        result = await _parse_resume_single(cleaned_text, request_id)
        mode = "single"

    await _emit_progress(
        progress_callback,
        {"type": "progress", "node": "merge_result", "title": "整理解析结果"},
    )
    total_ms = (perf_counter() - total_start) * 1000
    _log_timing(
        "parse_resume_total",
        total_ms,
        request_id,
        {"mode": mode, "input_length": len(cleaned_text)},
    )
    return _normalize_parse_result(result)


async def _parse_resume_from_attachment(
    *,
    encoded_file: str,
    filename: str,
    file_mime_type: str,
    request_id: Optional[str],
    progress_callback: ParseProgressCallback = None,
) -> Dict[str, Any]:
    await _emit_progress(
        progress_callback,
        {"type": "progress", "node": "request_ai", "title": "调用 AI 解析附件"},
    )
    total_start = perf_counter()
    messages = _build_resume_attachment_messages(
        RESUME_PARSING_PROMPT,
        filename,
        file_mime_type,
        encoded_file,
    )
    result = await _call_resume_llm(
        messages,
        request_id,
        "ai_call",
        {
            "filename": filename,
            "content_type": file_mime_type,
            "mode": "attachment",
        },
    )
    await _emit_progress(
        progress_callback,
        {"type": "progress", "node": "merge_result", "title": "整理解析结果"},
    )
    total_ms = (perf_counter() - total_start) * 1000
    _log_timing(
        "parse_resume_total",
        total_ms,
        request_id,
        {"mode": "attachment"},
    )
    return _normalize_parse_result(result)


async def parse_resume(
    file_data: bytes,
    filename: str,
    file_mime_type: str,
    request_id: Optional[str] = None,
    progress_callback: ParseProgressCallback = None,
) -> Dict[str, Any]:
    if not file_data:
        raise ValueError("文件为空，无法解析。")
    _ensure_attachment_size_limit(file_data)

    await _emit_progress(
        progress_callback,
        {"type": "progress", "node": "extract_text", "title": "提取简历正文"},
    )
    cleaned = _prepare_resume_text(
        extract_resume_text(
            file_data=file_data,
            filename=filename,
            file_mime_type=file_mime_type,
            request_id=request_id,
        ),
        request_id,
    )

    if not cleaned:
        encode_start = perf_counter()
        encoded_file = _encode_upload_data(file_data)
        encode_ms = (perf_counter() - encode_start) * 1000
        _log_timing(
            "encode_file",
            encode_ms,
            request_id,
            {"encoded_length": len(encoded_file)},
        )
        return await _parse_resume_from_attachment(
            encoded_file=encoded_file,
            filename=filename,
            file_mime_type=file_mime_type,
            request_id=request_id,
            progress_callback=progress_callback,
        )

    return await _parse_resume_from_text(
        cleaned_text=cleaned,
        request_id=request_id,
        progress_callback=progress_callback,
    )


async def parse_resume_with_thoughts(
    file_data: bytes,
    filename: str,
    file_mime_type: str,
    request_id: Optional[str] = None,
    progress_callback: ParseProgressCallback = None,
    thought_callback: ThoughtCallback = None,
) -> Dict[str, Any]:
    if not file_data:
        raise ValueError("文件为空，无法解析。")
    _ensure_attachment_size_limit(file_data)

    await _emit_progress(
        progress_callback,
        {"type": "progress", "node": "extract_text", "title": "提取简历正文"},
    )
    cleaned = _prepare_resume_text(
        extract_resume_text(
            file_data=file_data,
            filename=filename,
            file_mime_type=file_mime_type,
            request_id=request_id,
        ),
        request_id,
    )

    if not cleaned:
        encode_start = perf_counter()
        encoded_file = _encode_upload_data(file_data)
        encode_ms = (perf_counter() - encode_start) * 1000
        _log_timing(
            "encode_file",
            encode_ms,
            request_id,
            {"encoded_length": len(encoded_file)},
        )
        return await _parse_resume_from_attachment(
            encoded_file=encoded_file,
            filename=filename,
            file_mime_type=file_mime_type,
            request_id=request_id,
            progress_callback=progress_callback,
        )

    if not settings.gemini_api_key:
        logger.warning(
            "[ResumeParse] GEMINI_API_KEY missing, fallback to standard parser request_id=%s",
            request_id,
        )
        return await _parse_resume_from_text(
            cleaned_text=cleaned,
            request_id=request_id,
            progress_callback=progress_callback,
        )

    await _emit_progress(
        progress_callback,
        {"type": "progress", "node": "request_ai", "title": "调用 Gemini Thinking 解析"},
    )
    try:
        result = await _stream_resume_thinking_parse(
            cleaned_text=cleaned,
            request_id=request_id,
            thought_callback=thought_callback,
        )
    except Exception as exc:
        logger.warning(
            "[ResumeParse] Gemini Thinking unavailable, fallback to standard parser request_id=%s error=%s",
            request_id,
            str(exc),
        )
        await _emit_thought(thought_callback, {"type": "thought_reset"})
        return await _parse_resume_from_text(
            cleaned_text=cleaned,
            request_id=request_id,
            progress_callback=progress_callback,
        )
    await _emit_progress(
        progress_callback,
        {"type": "progress", "node": "merge_result", "title": "整理结构化结果"},
    )
    return result


async def fetch_existing_experiences(
    session: AsyncSession, user_id: str
) -> List[ExistingExperience]:
    results: List[ExistingExperience] = []
    offset = 0
    while True:
        batch = await _list_active_experiences(session, user_id, MAX_LIMIT, offset)
        if not batch:
            break
        for master, version in batch:
            if not version:
                continue
            results.append(
                ExistingExperience(
                    category=master.category,
                    title=version.title,
                    org=version.org or "",
                )
            )
        if len(batch) < MAX_LIMIT:
            break
        offset += MAX_LIMIT
    return results


async def _list_active_experiences(
    session: AsyncSession, user_id: str, limit: int, offset: int
) -> List[Tuple[MasterExperience, Optional[ExperienceVersion]]]:
    statement = (
        select(MasterExperience, ExperienceVersion)
        .join(
            ExperienceVersion,
            ExperienceVersion.id == MasterExperience.latest_version_id,
            isouter=True,
        )
        .where(
            MasterExperience.user_id == user_id,
            MasterExperience.is_archived.is_(False),
        )
        .order_by(desc(MasterExperience.updated_at))
        .limit(limit)
        .offset(offset)
    )
    result = await session.execute(statement)
    return list(result.all())


def _build_duplicate_index(
    entries: Iterable[ExistingExperience],
) -> Dict[ExperienceCategory, Tuple[List[str], set]]:
    index: Dict[ExperienceCategory, Tuple[List[str], set]] = {}
    for entry in entries:
        signature = _build_signature(entry.title, entry.org)
        if not signature:
            continue
        bucket = index.get(entry.category)
        if not bucket:
            bucket = ([], set())
            index[entry.category] = bucket
        bucket[0].append(signature)
        bucket[1].add(signature)
    return index


def _find_duplicate(
    signature: str, bucket: Optional[Tuple[List[str], set]]
) -> DuplicateMatch:
    if not signature or not bucket:
        return DuplicateMatch(is_duplicate=False)
    signatures, signature_set = bucket
    if signature in signature_set:
        return DuplicateMatch(is_duplicate=True, match_type="exact", match_score=1.0)
    best_score = 0.0
    for existing in signatures:
        score = _similarity(signature, existing)
        if score > best_score:
            best_score = score
    if best_score >= DUPLICATE_SIMILARITY_THRESHOLD:
        return DuplicateMatch(
            is_duplicate=True,
            match_type="similar",
            match_score=round(best_score, 2),
        )
    return DuplicateMatch(is_duplicate=False)


def apply_duplicate_flags(
    items: List[ParsedExperienceItem],
    existing_entries: Iterable[ExistingExperience],
) -> List[ParsedExperienceItem]:
    # Duplicates are detected per category by normalized (org + title) signature,
    # falling back to similarity when exact matches are not found.
    index = _build_duplicate_index(existing_entries)
    for item in items:
        signature = _build_signature(item.version.title, item.version.org)
        item.duplicate = _find_duplicate(signature, index.get(item.category))
    return items
