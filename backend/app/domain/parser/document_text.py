from __future__ import annotations

import io
import logging
import unicodedata
from pathlib import Path
from time import perf_counter
from typing import Optional
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from fastapi import UploadFile
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from .payload_normalization import _clean_resume_text

logger = logging.getLogger(__name__)

SUPPORTED_PDF_TYPES = {"application/pdf"}
SUPPORTED_DOCX_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
MIN_TEXT_LENGTH = 30
MIN_MEANINGFUL_TEXT_CHARS = 10
MAX_RESUME_TEXT_CHARS = 12_000
MAX_ATTACHMENT_BYTES = 5 * 1024 * 1024
UNREADABLE_RESUME_TEXT_ERROR = (
    "无法读取附件中的文本内容，请检查上传内容；当前不支持无法读取文本的附件。"
)

_DOCUMENT_LOG_WARN_THRESHOLDS_MS = {
    "read_file": 3_000,
    "parse_pdf": 8_000,
    "parse_docx": 5_000,
}


def _log_document_timing(
    step: str,
    duration_ms: float,
    request_id: Optional[str],
    extra: Optional[dict] = None,
) -> None:
    payload = {"step": step, "duration_ms": round(duration_ms, 2)}
    if request_id:
        payload["request_id"] = request_id
    if extra:
        payload.update(extra)
    threshold = _DOCUMENT_LOG_WARN_THRESHOLDS_MS.get(step)
    if threshold is not None and duration_ms >= threshold:
        logger.warning("[ResumeParse] %s", payload)
        return
    logger.info("[ResumeParse] %s", payload)


def _resolve_file_kind(file: UploadFile) -> str:
    filename = file.filename or ""
    extension = Path(filename).suffix.lower()
    content_type = (file.content_type or "").lower()
    if content_type in SUPPORTED_PDF_TYPES or extension == ".pdf":
        return "pdf"
    if content_type in SUPPORTED_DOCX_TYPES or extension == ".docx":
        return "docx"
    raise ValueError("不支持的文件类型，请上传 PDF 或 DOCX 文件。")


def _ensure_attachment_size_limit(data: bytes) -> None:
    if len(data) <= MAX_ATTACHMENT_BYTES:
        return
    max_mb = MAX_ATTACHMENT_BYTES / (1024 * 1024)
    raise ValueError(
        f"文件过大，无法直接解析。请上传不超过 {max_mb:.0f}MB 的 PDF 或 DOCX 文件。"
    )


def _resolve_file_mime(file: UploadFile, kind: str) -> str:
    content_type = (file.content_type or "").lower().strip()
    if content_type:
        return content_type
    if kind == "pdf":
        return "application/pdf"
    return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _resolve_file_kind_from_metadata(filename: str, content_type: str) -> str:
    extension = Path(filename or "").suffix.lower()
    normalized_type = (content_type or "").lower().strip()
    if normalized_type in SUPPORTED_PDF_TYPES or extension == ".pdf":
        return "pdf"
    if normalized_type in SUPPORTED_DOCX_TYPES or extension == ".docx":
        return "docx"
    raise ValueError("不支持的文件类型，请上传 PDF 或 DOCX 文件。")


async def extract_text(file: UploadFile, request_id: Optional[str] = None) -> bytes:
    total_start = perf_counter()
    read_start = perf_counter()
    data = await file.read()
    read_ms = (perf_counter() - read_start) * 1000
    _log_document_timing(
        "read_file",
        read_ms,
        request_id,
        {
            "size": len(data),
            "filename": file.filename or "",
            "content_type": file.content_type or "",
        },
    )
    if not data:
        raise ValueError("文件为空，无法解析。")
    _resolve_file_kind(file)
    total_ms = (perf_counter() - total_start) * 1000
    _log_document_timing("extract_text_total", total_ms, request_id)
    return data


def _extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def _extract_docx_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

    def append_table_text(table) -> None:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_parts = [para.text.strip() for para in cell.paragraphs if para.text.strip()]
                for nested_table in cell.tables:
                    append_table_text(nested_table)
                if cell_parts:
                    cells.append("\n".join(cell_parts))
            if cells:
                parts.append(" | ".join(cells))

    for table in doc.tables:
        append_table_text(table)

    return "\n".join(parts)


def extract_resume_text(
    file_data: bytes,
    filename: str,
    file_mime_type: str,
    request_id: Optional[str] = None,
) -> str:
    if not file_data:
        raise ValueError("文件为空，无法解析。")
    _ensure_attachment_size_limit(file_data)
    kind = _resolve_file_kind_from_metadata(filename, file_mime_type)
    parse_start = perf_counter()
    try:
        if kind == "pdf":
            text = _extract_pdf_text(file_data)
            parse_step = "parse_pdf"
        else:
            text = _extract_docx_text(file_data)
            parse_step = "parse_docx"
    except (PdfReadError, BadZipFile, PackageNotFoundError, ValueError) as exc:
        logger.warning(
            "[ResumeParse] failed to extract resume text request_id=%s filename=%s kind=%s error=%s",
            request_id,
            filename,
            kind,
            str(exc),
        )
        raise ValueError("文件无法读取，请确认文件未损坏、未加密且内容可解析。") from exc
    parse_ms = (perf_counter() - parse_start) * 1000
    _log_document_timing(parse_step, parse_ms, request_id, {"text_length": len(text)})
    return text


def _count_meaningful_text_chars(text: str) -> int:
    return sum(
        1
        for char in text
        if unicodedata.category(char).startswith(("L", "N"))
    )


def _prepare_resume_text(text: str, request_id: Optional[str] = None) -> Optional[str]:
    cleaned = _clean_resume_text(text)
    stripped = cleaned.strip()
    meaningful_char_count = _count_meaningful_text_chars(stripped)
    if meaningful_char_count < MIN_MEANINGFUL_TEXT_CHARS:
        logger.info(
            "[ResumeParse] extracted text unreadable, rejecting parse request_id=%s text_length=%s meaningful_char_count=%s",
            request_id,
            len(stripped),
            meaningful_char_count,
        )
        return None
    if len(stripped) < MIN_TEXT_LENGTH:
        logger.info(
            "[ResumeParse] extracted text too short, rejecting parse request_id=%s text_length=%s meaningful_char_count=%s",
            request_id,
            len(stripped),
            meaningful_char_count,
        )
        return None
    return cleaned
