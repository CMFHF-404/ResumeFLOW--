from __future__ import annotations

import asyncio
import ctypes
import io
import multiprocessing
import re
import threading
import sys
from pathlib import PurePosixPath
from typing import Callable
from zipfile import BadZipFile, ZIP_DEFLATED, ZIP_STORED, ZipFile, ZipInfo

from docx import Document
from fastapi import HTTPException
from pypdf import PdfReader


MAX_DOCUMENT_INPUT_BYTES = 5 * 1024 * 1024
MAX_DOCX_ENTRY_COUNT = 2_048
MAX_DOCX_ENTRY_UNCOMPRESSED_BYTES = 8 * 1024 * 1024
MAX_DOCX_TOTAL_UNCOMPRESSED_BYTES = 32 * 1024 * 1024
MAX_DOCX_XML_BYTES = 16 * 1024 * 1024
MAX_DOCX_XML_NODES = 100_000
MAX_DOCX_COMPRESSION_RATIO = 200.0
DOCX_RATIO_CHECK_MIN_BYTES = 1 * 1024 * 1024
MAX_PDF_PAGES = 200
MAX_PDF_OBJECTS = 10_000
MAX_PDF_STREAMS = 512
MAX_PDF_STREAM_UNCOMPRESSED_BYTES = 8 * 1024 * 1024
DOCUMENT_PARSE_CONCURRENCY = 2
DOCUMENT_PARSE_PENDING_LIMIT = 8
DOCUMENT_PARSE_TIMEOUT_SECONDS = 15.0
DOCUMENT_PARSE_PROCESS_MEMORY_BYTES = 384 * 1024 * 1024
DOCUMENT_PARSE_PROCESS_CPU_SECONDS = 20

_PDF_OBJECT_PATTERN = re.compile(rb"(?m)^\s*\d+\s+\d+\s+obj\b")
_PDF_PAGE_PATTERN = re.compile(rb"/Type\s*/Page(?!s)\b")
_PDF_STREAM_PATTERN = re.compile(rb"(?:\r\n|\n|\r)stream(?:\r\n|\n|\r)")
_PDF_UNBOUNDED_FILTER_PATTERN = re.compile(
    rb"/(?:LZWDecode|LZW|RunLengthDecode|RL)(?![A-Za-z])"
)
_ALLOWED_DOCX_COMPRESSION_TYPES = {ZIP_STORED, ZIP_DEFLATED}

_document_parse_gate: asyncio.Semaphore | None = None
_document_parse_gate_loop: asyncio.AbstractEventLoop | None = None
_pending_lock = threading.Lock()
_pending_document_parses = 0


class DocumentExtractionLimitError(ValueError):
    pass


class DocumentExtractionInvalidError(ValueError):
    pass


class DocumentExtractionBusyError(HTTPException):
    def __init__(self) -> None:
        super().__init__(
            status_code=503,
            detail="文档解析服务繁忙，请稍后重试。",
        )


class DocumentExtractionTimeoutError(HTTPException):
    def __init__(self) -> None:
        super().__init__(
            status_code=504,
            detail="文档解析超时，请稍后重试。",
        )


def _ensure_input_budget(data: bytes) -> None:
    if not data:
        raise DocumentExtractionInvalidError("文档为空，无法解析。")
    if len(data) > MAX_DOCUMENT_INPUT_BYTES:
        raise DocumentExtractionLimitError("文档超过安全处理大小限制。")


def _validate_docx_path(info: ZipInfo) -> None:
    raw_name = info.filename
    normalized_name = raw_name.replace("\\", "/")
    path = PurePosixPath(normalized_name)
    if (
        not normalized_name
        or "\x00" in normalized_name
        or normalized_name.startswith("/")
        or path.is_absolute()
        or any(part in {"", ".", ".."} for part in path.parts)
        or (path.parts and ":" in path.parts[0])
    ):
        raise DocumentExtractionLimitError("DOCX 包含不安全的归档路径。")


def _validate_docx_archive(data: bytes) -> None:
    try:
        with ZipFile(io.BytesIO(data)) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_DOCX_ENTRY_COUNT:
                raise DocumentExtractionLimitError("DOCX 归档条目过多。")

            total_uncompressed = 0
            total_xml_bytes = 0
            total_xml_nodes = 0
            for info in entries:
                _validate_docx_path(info)
                if info.flag_bits & 0x1:
                    raise DocumentExtractionLimitError("DOCX 不支持加密归档条目。")
                if info.compress_type not in _ALLOWED_DOCX_COMPRESSION_TYPES:
                    raise DocumentExtractionLimitError("DOCX 使用了不受支持的压缩格式。")
                if info.is_dir():
                    continue
                if info.file_size > MAX_DOCX_ENTRY_UNCOMPRESSED_BYTES:
                    raise DocumentExtractionLimitError("DOCX 单个归档条目展开后过大。")
                total_uncompressed += info.file_size
                if total_uncompressed > MAX_DOCX_TOTAL_UNCOMPRESSED_BYTES:
                    raise DocumentExtractionLimitError("DOCX 归档展开后总大小过大。")
                if info.file_size >= DOCX_RATIO_CHECK_MIN_BYTES:
                    ratio = info.file_size / max(info.compress_size, 1)
                    if ratio > MAX_DOCX_COMPRESSION_RATIO:
                        raise DocumentExtractionLimitError("DOCX 压缩比超过安全限制。")

                lowered_name = info.filename.casefold()
                if not lowered_name.endswith((".xml", ".rels")):
                    continue
                with archive.open(info) as entry:
                    xml_payload = entry.read(MAX_DOCX_ENTRY_UNCOMPRESSED_BYTES + 1)
                if len(xml_payload) > MAX_DOCX_ENTRY_UNCOMPRESSED_BYTES:
                    raise DocumentExtractionLimitError("DOCX XML 条目展开后过大。")
                total_xml_bytes += len(xml_payload)
                if total_xml_bytes > MAX_DOCX_XML_BYTES:
                    raise DocumentExtractionLimitError("DOCX XML 展开后总大小过大。")
                total_xml_nodes += xml_payload.count(b"<")
                if total_xml_nodes > MAX_DOCX_XML_NODES:
                    raise DocumentExtractionLimitError("DOCX XML 节点数量过多。")
    except BadZipFile as exc:
        raise DocumentExtractionInvalidError("DOCX 归档无效或已损坏。") from exc


def _count_with_limit(pattern: re.Pattern[bytes], data: bytes, limit: int, message: str) -> None:
    count = 0
    for _match in pattern.finditer(data):
        count += 1
        if count > limit:
            raise DocumentExtractionLimitError(message)


def _validate_pdf_structure(data: bytes) -> None:
    if not data.lstrip().startswith(b"%PDF-"):
        raise DocumentExtractionInvalidError("PDF 文件头无效。")
    _count_with_limit(
        _PDF_OBJECT_PATTERN,
        data,
        MAX_PDF_OBJECTS,
        "PDF 对象数量过多。",
    )
    _count_with_limit(
        _PDF_PAGE_PATTERN,
        data,
        MAX_PDF_PAGES,
        "PDF 页数超过安全限制。",
    )
    _count_with_limit(
        _PDF_STREAM_PATTERN,
        data,
        MAX_PDF_STREAMS,
        "PDF stream 数量过多。",
    )
    if _PDF_UNBOUNDED_FILTER_PATTERN.search(data):
        raise DocumentExtractionLimitError(
            "PDF 使用了无法安全限制展开大小的压缩过滤器。"
        )


def extract_pdf_text(data: bytes, *, max_output_chars: int) -> str:
    reader = PdfReader(io.BytesIO(data))
    if reader.is_encrypted:
        raise DocumentExtractionInvalidError("不支持加密 PDF。")
    if len(reader.pages) > MAX_PDF_PAGES:
        raise DocumentExtractionLimitError("PDF 页数超过安全限制。")
    parts: list[str] = []
    total_chars = 0
    for page in reader.pages:
        page_text = page.extract_text() or ""
        total_chars += len(page_text)
        if total_chars > max_output_chars:
            raise DocumentExtractionLimitError("文档提取文本超过安全字符限制。")
        parts.append(page_text)
    return "\n".join(parts)


def extract_docx_text(data: bytes, *, max_output_chars: int) -> str:
    document = Document(io.BytesIO(data))
    parts: list[str] = []
    total_chars = 0

    def append_text(value: str) -> None:
        nonlocal total_chars
        cleaned = value.strip()
        if not cleaned:
            return
        total_chars += len(cleaned)
        if total_chars > max_output_chars:
            raise DocumentExtractionLimitError("文档提取文本超过安全字符限制。")
        parts.append(cleaned)

    for paragraph in document.paragraphs:
        append_text(paragraph.text)

    def append_table_text(table) -> None:
        for row in table.rows:
            cells: list[str] = []
            for cell in row.cells:
                cell_parts = [
                    paragraph.text.strip()
                    for paragraph in cell.paragraphs
                    if paragraph.text.strip()
                ]
                for nested_table in cell.tables:
                    append_table_text(nested_table)
                if cell_parts:
                    cells.append("\n".join(cell_parts))
            if cells:
                append_text(" | ".join(cells))

    for table in document.tables:
        append_table_text(table)
    return "\n".join(parts)


DocumentExtractor = Callable[..., str]


def extract_document_text_sync(
    data: bytes,
    *,
    kind: str,
    max_output_chars: int,
    pdf_extractor: DocumentExtractor = extract_pdf_text,
    docx_extractor: DocumentExtractor = extract_docx_text,
) -> str:
    _ensure_input_budget(data)
    if max_output_chars <= 0:
        raise DocumentExtractionLimitError("文档提取文本字符限制无效。")
    if kind == "pdf":
        _validate_pdf_structure(data)
        text = pdf_extractor(data, max_output_chars=max_output_chars)
    elif kind == "docx":
        _validate_docx_archive(data)
        text = docx_extractor(data, max_output_chars=max_output_chars)
    else:
        raise DocumentExtractionInvalidError("不支持的文档类型。")
    if not isinstance(text, str):
        raise DocumentExtractionInvalidError("文档提取器返回了无效结果。")
    if len(text) > max_output_chars:
        raise DocumentExtractionLimitError("文档提取文本超过安全字符限制。")
    return text


_child_job_handle = None
_active_process_lock = threading.Lock()
_active_document_process_ids: set[int] = set()


def _apply_posix_process_limits() -> None:
    import resource

    memory_limit = DOCUMENT_PARSE_PROCESS_MEMORY_BYTES
    current_soft, current_hard = resource.getrlimit(resource.RLIMIT_AS)
    if current_hard not in (-1, resource.RLIM_INFINITY):
        memory_limit = min(memory_limit, current_hard)
    resource.setrlimit(resource.RLIMIT_AS, (memory_limit, memory_limit))
    cpu_limit = max(1, int(DOCUMENT_PARSE_PROCESS_CPU_SECONDS))
    resource.setrlimit(resource.RLIMIT_CPU, (cpu_limit, cpu_limit))


def _apply_windows_process_limits() -> None:
    from ctypes import wintypes

    class _IoCounters(ctypes.Structure):
        _fields_ = [
            ("ReadOperationCount", ctypes.c_ulonglong),
            ("WriteOperationCount", ctypes.c_ulonglong),
            ("OtherOperationCount", ctypes.c_ulonglong),
            ("ReadTransferCount", ctypes.c_ulonglong),
            ("WriteTransferCount", ctypes.c_ulonglong),
            ("OtherTransferCount", ctypes.c_ulonglong),
        ]

    class _BasicLimitInformation(ctypes.Structure):
        _fields_ = [
            ("PerProcessUserTimeLimit", ctypes.c_longlong),
            ("PerJobUserTimeLimit", ctypes.c_longlong),
            ("LimitFlags", wintypes.DWORD),
            ("MinimumWorkingSetSize", ctypes.c_size_t),
            ("MaximumWorkingSetSize", ctypes.c_size_t),
            ("ActiveProcessLimit", wintypes.DWORD),
            ("Affinity", ctypes.c_size_t),
            ("PriorityClass", wintypes.DWORD),
            ("SchedulingClass", wintypes.DWORD),
        ]

    class _ExtendedLimitInformation(ctypes.Structure):
        _fields_ = [
            ("BasicLimitInformation", _BasicLimitInformation),
            ("IoInfo", _IoCounters),
            ("ProcessMemoryLimit", ctypes.c_size_t),
            ("JobMemoryLimit", ctypes.c_size_t),
            ("PeakProcessMemoryUsed", ctypes.c_size_t),
            ("PeakJobMemoryUsed", ctypes.c_size_t),
        ]

    job_object_extended_limit_information = 9
    job_object_limit_process_time = 0x00000002
    job_object_limit_process_memory = 0x00000100
    job_object_limit_kill_on_job_close = 0x00002000
    kernel32 = ctypes.WinDLL("kernel32", use_last_error=True)
    kernel32.CreateJobObjectW.restype = wintypes.HANDLE
    kernel32.GetCurrentProcess.restype = wintypes.HANDLE
    kernel32.SetInformationJobObject.argtypes = [
        wintypes.HANDLE,
        ctypes.c_int,
        ctypes.c_void_p,
        wintypes.DWORD,
    ]
    kernel32.SetInformationJobObject.restype = wintypes.BOOL
    kernel32.AssignProcessToJobObject.argtypes = [
        wintypes.HANDLE,
        wintypes.HANDLE,
    ]
    kernel32.AssignProcessToJobObject.restype = wintypes.BOOL
    kernel32.CloseHandle.argtypes = [wintypes.HANDLE]
    kernel32.CloseHandle.restype = wintypes.BOOL
    job_handle = kernel32.CreateJobObjectW(None, None)
    if not job_handle:
        raise ctypes.WinError(ctypes.get_last_error())
    limits = _ExtendedLimitInformation()
    limits.BasicLimitInformation.PerProcessUserTimeLimit = int(
        DOCUMENT_PARSE_PROCESS_CPU_SECONDS * 10_000_000
    )
    limits.BasicLimitInformation.LimitFlags = (
        job_object_limit_process_time
        | job_object_limit_process_memory
        | job_object_limit_kill_on_job_close
    )
    limits.ProcessMemoryLimit = DOCUMENT_PARSE_PROCESS_MEMORY_BYTES
    configured = kernel32.SetInformationJobObject(
        job_handle,
        job_object_extended_limit_information,
        ctypes.byref(limits),
        ctypes.sizeof(limits),
    )
    assigned = configured and kernel32.AssignProcessToJobObject(
        job_handle,
        kernel32.GetCurrentProcess(),
    )
    if not assigned:
        error = ctypes.WinError(ctypes.get_last_error())
        kernel32.CloseHandle(job_handle)
        raise error
    global _child_job_handle
    _child_job_handle = job_handle


def _apply_child_process_limits() -> None:
    if sys.platform == "win32":
        _apply_windows_process_limits()
    else:
        _apply_posix_process_limits()
    import pypdf.filters as pdf_filters

    pdf_filters.ZLIB_MAX_OUTPUT_LENGTH = MAX_PDF_STREAM_UNCOMPRESSED_BYTES


def _document_process_worker(
    send_connection,
    data: bytes,
    kind: str,
    max_output_chars: int,
) -> None:
    try:
        try:
            _apply_child_process_limits()
        except BaseException:
            send_connection.send(
                ("unavailable", "文档解析隔离环境不可用，请稍后重试。")
            )
            return
        text = extract_document_text_sync(
            data,
            kind=kind,
            max_output_chars=max_output_chars,
        )
        send_connection.send(("ok", text))
    except DocumentExtractionLimitError as exc:
        send_connection.send(("limit", str(exc)))
    except DocumentExtractionInvalidError as exc:
        send_connection.send(("invalid", str(exc)))
    except BaseException:
        try:
            send_connection.send(("invalid", "文档无法安全解析。"))
        except BaseException:
            pass
    finally:
        send_connection.close()


def _start_document_process(
    data: bytes,
    kind: str,
    max_output_chars: int,
):
    context = multiprocessing.get_context("spawn")
    receive_connection, send_connection = context.Pipe(duplex=False)
    process = context.Process(
        target=_document_process_worker,
        args=(send_connection, data, kind, max_output_chars),
        daemon=True,
    )
    try:
        process.start()
    except BaseException:
        receive_connection.close()
        send_connection.close()
        raise
    send_connection.close()
    if process.pid is not None:
        with _active_process_lock:
            _active_document_process_ids.add(process.pid)
    return process, receive_connection


def _stop_and_join_document_process(process) -> None:
    process.join(timeout=0.2)
    if process.is_alive():
        process.terminate()
        process.join(timeout=1.0)
    if process.is_alive() and hasattr(process, "kill"):
        process.kill()
        process.join(timeout=1.0)
    if process.is_alive():
        raise RuntimeError("document extraction process could not be terminated")
    if process.pid is not None:
        with _active_process_lock:
            _active_document_process_ids.discard(process.pid)
    process.close()


def active_document_process_ids() -> set[int]:
    with _active_process_lock:
        return set(_active_document_process_ids)


def _decode_process_result(message) -> str:
    if not isinstance(message, tuple) or len(message) != 2:
        raise DocumentExtractionInvalidError("文档解析进程返回了无效结果。")
    status, payload = message
    if status == "ok" and isinstance(payload, str):
        return payload
    if status == "limit":
        raise DocumentExtractionLimitError(str(payload))
    if status == "invalid":
        raise DocumentExtractionInvalidError(str(payload))
    if status == "unavailable":
        raise DocumentExtractionBusyError()
    raise DocumentExtractionInvalidError("文档解析进程返回了无效结果。")


async def _extract_document_text_in_process(
    data: bytes,
    *,
    kind: str,
    max_output_chars: int,
) -> str:
    process = None
    receive_connection = None
    try:
        process, receive_connection = _start_document_process(
            data,
            kind,
            max_output_chars,
        )
        while True:
            if receive_connection.poll():
                return _decode_process_result(receive_connection.recv())
            if not process.is_alive():
                if receive_connection.poll():
                    return _decode_process_result(receive_connection.recv())
                raise DocumentExtractionInvalidError(
                    "文档解析进程未返回有效结果。"
                )
            await asyncio.sleep(0.01)
    finally:
        if receive_connection is not None:
            receive_connection.close()
        if process is not None:
            await asyncio.shield(
                asyncio.to_thread(_stop_and_join_document_process, process)
            )


def _get_document_parse_gate() -> asyncio.Semaphore:
    global _document_parse_gate, _document_parse_gate_loop
    loop = asyncio.get_running_loop()
    if _document_parse_gate is None or _document_parse_gate_loop is not loop:
        _document_parse_gate = asyncio.Semaphore(DOCUMENT_PARSE_CONCURRENCY)
        _document_parse_gate_loop = loop
    return _document_parse_gate


def _reserve_pending_parse() -> None:
    global _pending_document_parses
    with _pending_lock:
        if _pending_document_parses >= DOCUMENT_PARSE_PENDING_LIMIT:
            raise DocumentExtractionBusyError()
        _pending_document_parses += 1


def _release_pending_parse() -> None:
    global _pending_document_parses
    with _pending_lock:
        _pending_document_parses = max(0, _pending_document_parses - 1)


def _release_after_worker(
    worker: asyncio.Task[str],
    gate: asyncio.Semaphore,
) -> None:
    try:
        worker.exception()
    except BaseException:
        pass
    gate.release()
    _release_pending_parse()


async def extract_document_text_bounded(
    data: bytes,
    *,
    kind: str,
    max_output_chars: int,
    pdf_extractor: DocumentExtractor = extract_pdf_text,
    docx_extractor: DocumentExtractor = extract_docx_text,
) -> str:
    _ensure_input_budget(data)
    _reserve_pending_parse()
    gate = _get_document_parse_gate()
    acquired = False
    deferred_release = False
    worker: asyncio.Task[str] | None = None
    use_isolated_process = (
        pdf_extractor is extract_pdf_text
        and docx_extractor is extract_docx_text
    )
    try:
        try:
            async with asyncio.timeout(DOCUMENT_PARSE_TIMEOUT_SECONDS):
                await gate.acquire()
                acquired = True
                if use_isolated_process:
                    return await _extract_document_text_in_process(
                        data,
                        kind=kind,
                        max_output_chars=max_output_chars,
                    )
                worker = asyncio.create_task(
                    asyncio.to_thread(
                        extract_document_text_sync,
                        data,
                        kind=kind,
                        max_output_chars=max_output_chars,
                        pdf_extractor=pdf_extractor,
                        docx_extractor=docx_extractor,
                    )
                )
                return await asyncio.shield(worker)
        except TimeoutError as exc:
            if worker is not None and not worker.done():
                worker.add_done_callback(
                    lambda completed: _release_after_worker(completed, gate)
                )
                deferred_release = True
            raise DocumentExtractionTimeoutError() from exc
        except asyncio.CancelledError:
            if worker is not None and not worker.done():
                worker.add_done_callback(
                    lambda completed: _release_after_worker(completed, gate)
                )
                deferred_release = True
            raise
    finally:
        if not deferred_release:
            if acquired:
                gate.release()
            _release_pending_parse()


def _reset_document_parse_gate_for_tests() -> None:
    global _document_parse_gate, _document_parse_gate_loop, _pending_document_parses
    _document_parse_gate = None
    _document_parse_gate_loop = None
    with _pending_lock:
        _pending_document_parses = 0
